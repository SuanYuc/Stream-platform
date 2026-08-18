from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(r'C:\Users\acidrain\PycharmProjects\Nsy_Broadcasting_platform')
OUT = ROOT / 'docs' / 'competition_manual' / 'AI智能增强导播平台项目汇报技术说明书.docx'
SCREEN_DIR = ROOT / 'docs' / 'competition_manual' / 'screenshots'
OUT.parent.mkdir(parents=True, exist_ok=True)


def set_run_font(run, east='宋体', ascii_font='Times New Roman', size=None, bold=None, color=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_para_format(p, before=0, after=6, line=1.15, first_line=True):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def shade_paragraph(p, fill='EEF3F8'):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    set_para_format(p, after=after)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, east='黑体', ascii_font='Times New Roman', size=15, bold=True, color=(31, 78, 121))
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, east='黑体', ascii_font='Times New Roman', size=12.5, bold=True, color=(46, 92, 130))
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, east='黑体', ascii_font='Times New Roman', size=11.5, bold=True, color=(64, 64, 64))
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, east='宋体', ascii_font='Times New Roman', size=10, color=(90, 90, 90))
    return p


def add_figure(doc, image_name, caption):
    path = SCREEN_DIR / image_name
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(14.8))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'[此处插入图片：{image_name}]')
        set_run_font(r, size=10, color=(120, 120, 120))
    add_caption(doc, caption)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    shade_paragraph(p, 'F2F4F7')
    r = p.add_run(code.strip())
    set_run_font(r, east='Consolas', ascii_font='Consolas', size=9.2, color=(40, 40, 40))
    return p


def add_keypoint(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.7)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.4)
section.right_margin = Cm(1.8)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.0)

# Base styles
styles = doc.styles
styles['Normal'].font.name = 'Times New Roman'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
styles['Normal'].font.size = Pt(11)
for name in ('Heading 1', 'Heading 2', 'Heading 3'):
    styles[name].font.name = 'Times New Roman'
    styles[name]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# Header / footer
hdr = section.header.paragraphs[0]
hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = hdr.add_run('AI智能增强导播平台项目汇报技术说明书')
set_run_font(r, east='宋体', ascii_font='Times New Roman', size=9, color=(100, 100, 100))
ftr = section.footer.paragraphs[0]
ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ftr.add_run('Nsy Broadcasting Platform')
set_run_font(r, east='宋体', ascii_font='Times New Roman', size=9, color=(100, 100, 100))

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(100)
r = p.add_run('AI智能增强导播平台')
set_run_font(r, east='黑体', ascii_font='Times New Roman', size=26, bold=True, color=(31, 78, 121))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('项目汇报技术说明书')
set_run_font(r, east='黑体', ascii_font='Times New Roman', size=22, bold=True, color=(31, 78, 121))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run('面向直播制作、语义智能导播、实时智能滤镜、无限画布编排与云推流的一体化系统')
set_run_font(r, east='宋体', ascii_font='Times New Roman', size=12.5, color=(80, 80, 80))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run('参考“技术交底书”式写作结构整理')
set_run_font(r, east='宋体', ascii_font='Times New Roman', size=11, color=(110, 110, 110))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('注：参考文档仅用于格式与写作风格，不沿用其中与本项目无关的技术内容。')
set_run_font(r, east='宋体', ascii_font='Times New Roman', size=10.5, color=(140, 70, 70))
doc.add_page_break()

add_h1(doc, '一、项目名称')
add_body(doc, '本项目名称为“AI智能增强导播平台”，工程名称为 Nsy Broadcasting Platform。该系统是一套面向直播制作和轻量化演播应用的本地导播与云端推流一体化平台，重点解决多源画面采集、场景图层合成、智能导播推荐、实时智能滤镜、无限画布编排、音轨调音和云端直播发布等关键问题。')
add_body(doc, '从技术表达上看，本项目可以概括为一种基于 PyQt6 桌面工作站、ONNX/MediaPipe 智能视觉处理、FG-CLIP2 图文语义匹配、画布化场景组织和 Go 云端流媒体服务的智能导播系统。它既保留传统导播台中 Preview 与 Program 分离、场景切换、推流录制等基础流程，又围绕“智能化、轻量化、可扩展”进行了功能扩展。')

add_h1(doc, '二、所属技术领域')
add_body(doc, '本项目属于广播电视工程、网络直播制播、计算机视觉、人工智能辅助媒体制作和云端流媒体服务交叉领域。具体而言，项目涉及桌面端多媒体软件开发、实时视频采集与合成、音视频编码与推流、ONNX 模型部署、语义图文匹配、MediaPipe 人像与人脸处理、窗口音频采集隔离以及基于 Go 的云端流媒体服务管理。')
add_body(doc, '在应用场景上，本项目可用于校园活动直播、课程直播、融媒体演播、赛事导播、综艺节目轻量化制作、远程会议包装和小型直播间搭建。与只负责采集或播放的普通工具不同，本系统围绕完整导播流程构建，强调“输入源管理、场景组织、节目输出、智能增强、云端发布”的闭环能力。')

add_h1(doc, '三、背景技术（现有技术）')
add_h2(doc, '3-1 最接近的现有技术')
add_body(doc, '当前直播制作领域已经存在 OBS Studio、vMix、Wirecast 等成熟导播软件。这类软件能够完成摄像头、窗口、屏幕、网络流等多种输入源采集，也能够通过场景和来源列表进行画面合成，并支持 RTMP 推流和本地录制。它们为个人直播、课程录制、会议转播和专业导播提供了有效工具。')
add_body(doc, '同时，人工智能技术在媒体制作中的使用正在增加。例如，基于图像分割的人像抠像技术可以替代传统绿幕，基于人脸关键点的 AR 贴纸能够增强画面表现，基于风格迁移的 ONNX 模型可以实现实时画面风格化，基于 CLIP 类图文模型的语义匹配能够把自然语言描述与图像内容关联起来。这些技术为导播系统从“人工控制工具”向“智能辅助工作站”演进提供了条件。')
add_body(doc, '在云端分发方面，传统直播系统通常使用 RTMP 推流到云服务商直播平台，再由云服务商生成播放地址。对于中小型项目而言，直接接入云平台有一定配置成本，也不便于在项目内部控制转推、监控和播放地址生成。因此，使用一台轻量云服务器承接本地导播台推流，再通过流媒体服务对外提供 RTMP/HLS 播放，是一种更易展示和部署的方案。')

add_h2(doc, '3-2 需要解决的技术问题')
add_body(doc, '现有导播系统在复杂使用场景中仍存在若干不足。首先，传统场景和图层管理多以线性列表为主，用户在复杂节目包装中难以直观理解多个场景之间的关系，也难以把多个图层作为空间对象自由组合。其次，导播切换高度依赖人工判断，当场景数量较多时，导播员需要快速观察多个缩略图并选择目标画面，容易出现查找慢、误切换或错过最佳时机的问题。')
add_body(doc, '其次，智能视觉能力通常与导播管线割裂。许多工具把抠像、AR、风格化处理作为插件或外部后处理，无法真正以图层为单位进行独立控制。本项目需要让每个图层都能单独启用基础滤镜、ONNX 风格迁移、MediaPipe 虚拟背景和 AR 贴纸，并在合成器中统一输出。')
add_body(doc, '再次，音频采集在直播系统中经常被简化为“系统声”或“麦克风”两个入口，难以针对窗口图层进行精确音轨隔离。实际直播中，用户常常只希望采集某个窗口的声音，而不希望其他应用声音进入节目。因此，本项目需要将音频抽象为音轨，并允许用户在调音台中进行音量、静音、频段和监听控制。')
add_body(doc, '最后，本地推流与云端分发之间需要形成稳定链路。系统不仅要把节目输出编码为 RTMP，还应在云服务器上提供可管理的服务入口，能够生成播放地址、启动转推、检查健康状态，并尽量减少服务器二次编码压力。')

add_h1(doc, '四、为解决技术问题而采用的总体技术方案')
add_body(doc, '为解决上述问题，本项目采用“本地智能导播工作站 + 云端轻量流媒体服务”的总体方案。本地端以 PyQt6 为界面框架，以 OpenCV、NumPy 和 PyAV 为实时视频处理基础，以 ONNX Runtime 和 MediaPipe 为智能处理能力，以 AppState 维护场景、图层、音轨和输出配置。云端采用 Go 编写控制服务，配合 MediaMTX 和 FFmpeg 实现 RTMP 接收、HLS 输出和直播转推。')
add_body(doc, '系统整体采用模块化设计。场景与图层模块负责描述节目内容，采集模块负责从各类输入源获取最新帧，渲染模块负责按图层优先级合成画面，输出模块负责录制和推流，音频模块负责采集、处理和分发音频，语义导播模块负责根据中文描述推荐场景，智能滤镜模块负责图层级画面增强，无限画布模块负责空间化编排和写回导播台，云推流模块负责服务器端流媒体发布。')
add_body(doc, '本项目的关键设计原则是尽量复用原有导播管线。无论是智能导播、智能滤镜还是无限画布，最终都要回到“场景和图层”这一基础数据结构中。这样可以保证 Preview、Program、推流、录制、转场和紧急占位逻辑保持一致，避免为新功能另起一套不兼容的输出链路。')
add_figure(doc, '01_main_console.png', '图1 AI智能增强导播平台主界面与基础导播工作流示意')

add_h1(doc, '五、有益效果')
add_body(doc, '第一，本项目降低了轻量化直播制作门槛。用户可以在一个桌面软件中完成输入源添加、场景组织、图层编辑、预览、节目输出、录制和推流，不需要在多个软件之间反复切换。')
add_body(doc, '第二，本项目提升了导播选择效率。语义智能导播模块把用户输入的自然语言转换为场景推荐结果，导播员可以通过“主持人特写”“产品展示”“观众画面”等中文描述快速定位候选场景，从而减少人工查找时间。')
add_body(doc, '第三，本项目增强了画面表现力。智能滤镜被设计为图层级能力，用户可以对不同图层分别启用 ONNX 风格迁移、虚拟背景、AR 贴纸、色彩校正和马赛克，从而实现更灵活的画面包装。')
add_body(doc, '第四，本项目提升了复杂场景编排能力。无限画布模式把场景和图层从线性列表转换为空间对象，用户可以在画布中同时管理多个场景框和多个图层对象，并将组合结果写回导播台。')
add_body(doc, '第五，本项目增强了直播安全性。紧急占位机制允许用户在直播事故发生时快速将节目输出切换到占位场景，而编辑预览保持原状，便于后台继续修复正常场景。')
add_body(doc, '第六，本项目具备云端发布能力。通过 Go 服务、MediaMTX 和 FFmpeg，系统可以把本地节目流推送到阿里云服务器，再向观众提供 RTMP/HLS 播放地址，并为后续转推到云直播平台预留接口。')

add_h1(doc, '六、附图及附图说明')
add_body(doc, '图1 展示 AI智能增强导播平台主界面与基础导播工作流，包括场景管理、编辑预览、节目输出、推流录制和音轨选择。')
add_body(doc, '图2 展示图层与来源管理界面，重点体现输入源添加、图层优先级、滤镜、音频、色彩校正和智能增强入口。')
add_body(doc, '图3 展示语义智能导播界面，重点体现中文语义输入、FG-CLIP2 场景推荐、相似度分数和推理设备信息。')
add_body(doc, '图4 展示智能滤镜效果对比，包含虚拟背景、AR 贴纸、ONNX 卡通化、莫奈风格和梵高风格滤镜。')
add_body(doc, '图5 展示无限画布工作区，包含坐标网格、场景框、图层对象、折叠式资源列表和属性编辑入口。')
add_body(doc, '图6 展示云推流服务与播放验证，包含本地 RTMP 推流地址、Go API 服务状态、MediaMTX 播放地址和 HLS 验证结果。')

add_h1(doc, '七、具体实施方式')
add_h2(doc, '1. 导播平台基础功能的实现')
add_h3(doc, '1.1 场景、图层与状态管理')
add_body(doc, '本系统以 Scene 和 Layer 作为基础数据模型。Scene 表示一个导播场景，内部保存多个 Layer；Layer 表示一个可被合成到画布上的输入源对象，包含图层 ID、名称、类型、启用状态、锁定状态、坐标、尺寸、饱和度、对比度、色温、马赛克、音量、优先级和 source 字典。source 字典用于保存具体输入源参数，例如摄像头编号、窗口句柄、网络流地址、图片路径、视频路径、ONNX 风格参数、虚拟背景参数和 AR 参数。')
add_body(doc, 'AppState 作为全局状态中心，负责初始化场景、维护当前选中场景、保证紧急占位场景始终位于最后、管理图层增删改、归一化优先级、保存转场配置和音轨选择。系统默认生成九个普通场景，并单独生成紧急占位场景。普通场景可以被清空或调整数量，占位场景则被保护，避免用户误删导致直播事故时无可用兜底画面。')
add_body(doc, '图层遮挡关系由 priority 编号决定，编号越大，图层越靠上。渲染时，系统先绘制低编号图层，再绘制高编号图层。与传统列表拖拽相比，这种编号式优先级更明确，也便于在图层管理子界面和无限画布之间同步。')
add_code(doc, '''ordered_layers = sort(scene.layers, key=(layer.priority, original_index))
for layer in ordered_layers:
    if layer.enabled:
        render_layer_to_canvas(layer)''')

add_h3(doc, '1.2 输入源采集与复用')
add_body(doc, '采集模块位于 `nsy_broadcasting_platform/capture/` 目录。系统为摄像头、屏幕、窗口、网络流、图片和视频文件分别实现 Source 类。各类视频源继承或复用 BaseVideoSource 的线程模型，在独立线程中按设定 FPS 采集最新帧，并通过线程锁保存到缓存中。渲染线程只读取最新帧，避免帧队列堆积造成延迟。')
add_body(doc, 'SourceManager 负责根据场景中启用的图层同步采集源。系统会根据图层类型、输入源参数、采集质量、分辨率和帧率生成签名。当签名没有变化时直接复用旧 Source；当签名变化时停止旧 Source 并创建新 Source。这种设计减少了重复打开摄像头或网络流的次数，提高采集稳定性。')
add_code(doc, '''function sync_scenes(scenes):
    enabled_layers = collect_enabled_layers(scenes)
    keep_ids = set(layer.id for layer in enabled_layers)
    stop_sources_not_in(keep_ids)
    for layer in enabled_layers:
        signature = build_signature(layer.type, layer.source, quality, fps)
        if source_exists(layer.id) and signature_unchanged(layer.id, signature):
            continue
        rebuild_source(layer.id, layer)''')

add_h3(doc, '1.3 图层渲染与节目合成')
add_body(doc, 'Compositor 是画面合成的核心。每次渲染一个场景时，系统先创建黑色背景画布，再按图层优先级逐层处理。对于普通视频图层，系统先执行饱和度、对比度、色温和马赛克等基础滤镜；如果该图层启用 ONNX 风格迁移，则继续执行风格化处理；如果启用 AR 贴纸，则调用 MediaPipe FaceMesh 检测人脸关键点并叠加贴纸；如果启用虚拟背景，则调用 MediaPipe Selfie Segmentation 生成人像掩码并与背景合成。')
add_body(doc, '最终，系统会将处理后的图层帧缩放到图层尺寸，并根据图层坐标裁剪到画布范围内。对于带透明通道的 PNG 图层，系统使用 Alpha 混合；对于普通 RGB 图层，系统直接覆盖对应画布区域。合成结果作为 RenderResult 返回，包含最终画面、图层矩形信息和智能处理指标。')
add_code(doc, '''Result = Source_RGB * Alpha + Destination_RGB * (1 - Alpha)

function render_scene(scene):
    canvas = black_canvas(width, height)
    for layer in ordered_layers:
        frame = source_manager.get_frame(layer.id)
        frame = apply_layer_filters(frame, layer)
        frame = resize(frame, layer.width, layer.height)
        canvas = blend_or_copy(canvas, frame, layer.x, layer.y)
    return canvas''')

add_h3(doc, '1.4 Preview、Program、转场与紧急占位')
add_body(doc, 'RenderThread 持续运行在独立线程中，负责生成编辑预览和节目输出。编辑预览显示当前正在编辑的场景，节目输出显示真正对外输出的场景。一般情况下二者相同；当紧急占位启用时，Program 切换为占位场景，而 Preview 仍保持当前编辑场景。这样导播员可以在后台继续修复正常场景，同时让观众端看到安全占位画面。')
add_body(doc, '转场模块位于 `render/transitions.py`，支持硬切、叠化、划像、DVE 和自定义媒体转场。系统在正常场景切换时触发转场，而紧急占位切换不会触发转场，以保证事故处理尽可能立即生效。节目延迟通过时间戳队列实现，系统根据当前时间减去延迟时间选择对应历史帧，并在切换场景或转场时清空延迟缓存，以减少闪烁和割裂。')
add_code(doc, '''target_time = now - delay_ms / 1000
program_frame = latest_frame_before(target_time)

if emergency_placeholder_enabled:
    program_scene = placeholder_scene
    transition.cancel()
else:
    program_scene = active_scene''')

add_h3(doc, '1.5 编码、录制与推流')
add_body(doc, 'OutputManager 管理录制 Worker 和推流 Worker。RenderThread 生成的 Program Frame 会统一推送给 OutputManager，再由 OutputManager 分发给正在运行的录制线程和推流线程。因此，本地节目输出、录制文件和 RTMP 推流使用同一份节目帧，避免输出画面不一致。')
add_body(doc, 'EncoderWorker 使用 PyAV 创建视频流和音频流。视频编码器候选包括 h264_nvenc 和 libx264，用户可以选择 GPU、CPU 或自动模式。推流模式下系统使用低延迟参数，例如 zerolatency、bf=0 和较快 preset；录制模式则可以使用更偏画质的参数。音频使用 AAC 编码，并根据音频时间戳校正视频 PTS，减少音画不同步。')
add_code(doc, '''video_pts = max(time_elapsed * fps,
                audio_pts / sample_rate * fps,
                last_video_pts + 1)

function encoder_worker_loop():
    open_container(target)
    build_video_stream(codec_candidates)
    build_audio_stream(aac)
    while running:
        encode_latest_video_frame()
        drain_audio_queue_and_encode()''')
add_figure(doc, '02_scene_layer_control.png', '图2 图层与来源管理功能示意')

add_h2(doc, '2. 智能导播功能的实现')
add_body(doc, '智能导播模块位于 `nsy_broadcasting_platform/semantic_director/`。该模块使用本地 `fgclip2_semantic` 文件夹中的图像编码器、文本编码器、tokenizer 和 metadata。系统不需要联网即可完成语义推荐，符合本地导播工具对稳定性的要求。')
add_body(doc, '用户输入中文语义后，系统先通过 tokenizer 将文本转换为 input_ids 和 attention_mask，再送入 FG-CLIP2 文本编码器得到文本向量。与此同时，系统将渲染线程生成的场景缩略图 QImage 转换为 RGB ndarray，并按照模型要求进行 resize、normalize 和 patchify，送入图像编码器得到图像向量。')
add_body(doc, '由于文本向量和图像向量位于同一语义空间，系统可以使用点积计算相似度。所有场景按分数从高到低排序，最高分且超过阈值的场景作为推荐结果。推荐任务由后台线程执行，界面通过信号接收结果，避免主线程卡顿。当前系统采用推荐模式，用户可以选择定位或切换推荐场景。')
add_code(doc, '''image_emb = normalize(image_encoder(scene_thumbnail))
text_emb  = normalize(text_encoder(query))
score     = dot(image_emb, text_emb)

best_scene = argmax(score_i)
if best_score >= threshold:
    return best_scene
else:
    return no_recommendation''')
add_body(doc, '智能导播的实现重点不在于替代导播员，而在于提高导播员对复杂场景的检索效率。系统把 AI 的角色限定为“推荐候选”，最终是否切换节目输出仍由用户确认。这种设计适合真实直播工作流，因为它兼顾智能化和安全性。')
add_figure(doc, '03_semantic_director.png', '图3 语义智能导播推荐界面示意')

add_h2(doc, '3. 智能滤镜功能的实现')
add_h3(doc, '3.1 基础滤镜与图层级处理链')
add_body(doc, '基础滤镜由 `utils.py::apply_video_filters` 实现，主要包括饱和度、对比度、色温和马赛克。饱和度调节通常通过 HSV 或等效色彩变换实现，对比度通过线性缩放增强或降低画面明暗差异，色温通过改变红蓝通道权重实现偏暖或偏冷，马赛克通过降采样再放大实现局部像素块化。')
add_body(doc, '所有滤镜都以图层为单位执行。也就是说，同一场景内的摄像头图层可以启用虚拟背景，窗口图层可以保持原样，图片图层可以启用风格迁移。这种图层级处理方式比全局后处理更灵活。')

add_h3(doc, '3.2 ONNX 风格迁移滤镜')
add_body(doc, 'ONNX 风格迁移滤镜由 `OnnxStyleTransferFilter` 实现。系统将 `onnx_models/cartoon.onnx` 作为卡通化模型，将 `monet.onnx` 作为莫奈风格模型，将 `vangogh.onnx` 作为梵高风格模型。用户在图层滤镜界面选择风格并点击应用后，图层 source 中会记录 `onnx_style` 参数，合成器在每帧渲染时读取该参数并调用对应滤镜。')
add_body(doc, '模型加载采用懒加载和缓存机制。用户第一次启用某个风格时，系统创建 ONNX session 并缓存；后续相同风格直接复用缓存。`gpu_runtime.create_session()` 会按 CUDA、DirectML、CPU 的顺序尝试创建会话，并启用 ONNX 图优化。如果模型不可用、输出无法解析或输出退化，系统会自动降级为本地近似滤镜，而不是让整个导播流程中断。')
add_code(doc, '''function apply_onnx_style_filter(frame, style):
    if style == none:
        return frame
    filter = style_cache.get(style)
    if filter is null:
        preload_filter_async(style)
        return fallback_style(frame, style)
    output = filter.run(frame)
    if output_invalid(output):
        output = fallback_style(frame, style)
    return blend(frame, output, strength)''')

add_h3(doc, '3.3 MediaPipe 虚拟背景')
add_body(doc, '虚拟背景功能由 `VirtualBackgroundFilter` 实现，底层使用 MediaPipe Selfie Segmentation。系统先把视频帧送入分割模型得到人像掩码，再对掩码进行时间平滑、阈值处理、腐蚀和模糊，使人物边缘更加稳定。随后系统根据用户选择，将人物前景与背景图片或模糊背景进行合成。')
add_code(doc, '''mask = mediapipe_segmentation(frame)
mask = temporal_smooth(mask)
soft_mask = erode_and_blur(mask)
result = foreground * soft_mask + background * (1 - soft_mask)''')

add_h3(doc, '3.4 MediaPipe AR 贴纸')
add_body(doc, 'AR 贴纸功能由 `FaceEffectFilter` 实现，底层使用 MediaPipe FaceMesh。系统通过鼻尖、双眼、额头和下巴等关键点估计人脸中心、尺度和角度，再将贴纸图片缩放、旋转并叠加到对应位置。项目内置狗鼻子、猫耳朵和卡通眼睛等素材，也支持后续扩展新的贴纸类型。')
add_body(doc, '为了减少贴纸抖动，系统对中心点、角度和尺度进行平滑处理。当短时间未检测到人脸时，可以保持上一帧贴纸状态，避免贴纸瞬间消失造成观感突兀。')
add_code(doc, '''landmarks = face_mesh(frame)
center = smooth(landmark[nose])
scale = distance(left_eye, right_eye)
angle = estimate_face_angle(forehead, chin)
sticker = resize_rotate(sticker, scale, angle)
frame = alpha_blend(frame, sticker, center)''')
add_figure(doc, '04_smart_filters.png', '图4 智能滤镜效果对比示意')

add_h2(doc, '4. 无限画布功能的实现')
add_body(doc, '无限画布模块位于 `nsy_broadcasting_platform/canvas/`。它的核心目标是把场景和图层从传统线性列表中释放出来，转化为可自由摆放、缩放和组合的空间对象。画布采用 PyQt6 Graphics View 体系实现，背景保留低对比坐标网格，用户可以在画布中进行平移、缩放、选择和拖拽。')
add_body(doc, '画布数据由 `CanvasDocument` 保存。该结构包含 document_id、name、version、viewport、output_frame、items、groups、history_metadata、created_at 和 updated_at。单个图层或场景框由 `CanvasItemModel` 表示，其中保存 item_id、type、source_ref、scene_ref、parent_item_id、位置、尺寸、旋转、透明度、可见性、锁定状态、z_index、裁剪、滤镜、抠像、音频和 metadata。')
add_body(doc, '在画布中，场景表现为一个场景框。场景框在逻辑上是图层集合，可以来源于导播台已有场景，也可以作为新建场景存在。图层是基本元素，可以独立存在，也可以被拖入某个场景框。图层拖入场景框时，系统视为将该图层加入对应场景；图层拖出场景框时，系统视为从该场景移出。这个过程只改变归属关系，不破坏图层自身的滤镜、音频、尺寸和来源参数。')
add_body(doc, '画布与导播台之间通过 `SceneCanvasAdapter` 和 `DirectorCanvasBridge` 连接。前者负责把导播台 Scene/Layer 转换为 CanvasDocument/CanvasItemModel，也负责把画布对象转回 Layer；后者负责把画布文档写回已有场景，或新建一个场景保存。由于最终仍写回 AppState 中的标准场景图层结构，所以原有渲染、推流、录制和转场逻辑无需重写。')
add_code(doc, '''function export_document_to_scene(document, target_scene):
    layers = document_to_layers(document)
    state.clear_scene_layers(target_scene)
    for layer in layers:
        state.add_layer(layer, scene_id=target_scene)
    refresh_scene_and_layer_ui()
    return success''')
add_body(doc, '无限画布还与智能滤镜和音频属性保持兼容。画布对象的 filters 字段保存饱和度、对比度、色温、马赛克和 ONNX 风格；chroma_key 字段保存虚拟背景和 AR 参数；audio 字段保存音量、静音、幅度和低中高频参数。用户在画布中打开图层特效子面板时，修改的参数会通过桥接层同步回导播台图层。')
add_figure(doc, '05_infinite_canvas.png', '图5 无限画布工作区与场景框编排示意')

add_h2(doc, '5. 音轨调音台与音频处理的实现')
add_body(doc, '音频模块采用音轨化设计。系统将音频来源抽象为自动音轨、系统声音、窗口音轨、麦克风和总音轨。每个窗口图层可以派生出一个窗口音轨，用户也可以选择系统声音或麦克风作为节目音频来源。')
add_body(doc, 'LoopbackCapture 负责采集系统输出声音，InputCapture 负责采集麦克风输入，SessionMatcher 负责根据进程 ID 或进程名匹配窗口音频会话。当用户选择单个窗口音轨并启用严格隔离时，系统会尝试屏蔽其他音频会话，使节目中只保留目标窗口声音。')
add_body(doc, 'AudioController 是音频处理中心。它会根据当前音轨配置选择采集方式，并对音频块执行静音、增益、幅度和低中高频处理。调音台界面采用数字调音台式 Channel Strip，每条音轨都有电平显示、音量推子、静音开关和频段控制。系统还预留了 AI Processor Hook，后续可以接入 ONNX 降噪、违禁词检测、自动增益或声纹识别模型。')
add_code(doc, '''function process_audio_chunk(chunk):
    if muted:
        return silence(chunk)
    samples = bytes_to_int16(chunk)
    samples = samples * volume * amplitude
    samples = apply_fft_eq(samples, low_gain, mid_gain, high_gain)
    for processor in ai_processors:
        samples = processor.process(samples)
    return int16_to_bytes(clamp(samples))''')

add_h2(doc, '6. GPU 适配与性能治理')
add_body(doc, '本项目涉及多类实时计算，包括视频采集、图层合成、智能滤镜、语义推荐、音频处理和编码推流。为保证运行流畅，系统在多个层面进行了性能治理。')
add_body(doc, '在模型推理层，`gpu_runtime.py` 会优先尝试 CUDAExecutionProvider、DmlExecutionProvider 和 TensorrtExecutionProvider 等 GPU Provider，再回退到 CPUExecutionProvider。ONNX session 创建时启用图优化，并限制合理线程数量，避免模型推理占满 CPU。')
add_body(doc, '在采集层，每个输入源使用独立线程保存最新帧，渲染线程只读取最新结果，避免旧帧堆积。在渲染层，图层智能滤镜实例被缓存，MediaPipe 和 ONNX 支持后台预热，避免用户首次点击时长时间阻塞。在输出层，推流编码优先使用 NVENC，失败后回退 libx264。云端转推优先复制视频流，减少二次编码压力。')
add_body(doc, '在 UI 层，音频电平动画只刷新自身组件，无限画布拖拽主要更新对象模型而非立刻重建整个导播场景。只有当用户写回场景或同步场景框内容时，画布才转换为导播台图层结构。这种设计可以减少编辑过程中的卡顿。')

add_h2(doc, '7. 云推流功能的实现')
add_body(doc, '云推流模块由本地 Python 导播台和云端 Go 服务共同完成。本地端将 Program Frame 交给 EncoderWorker 编码为 RTMP 流，并推送到云服务器。云服务器上运行 MediaMTX 接收 RTMP，并对外提供 RTMP 和 HLS 播放。Go 服务负责提供 API，生成推流地址和播放地址，并管理 FFmpeg Relay。')
add_body(doc, '当前默认推流地址为 `rtmp://<YOUR_SERVER_IP>:1935/live/main`。用户在本地导播台点击开始推流后，节目输出帧和音频块会被编码并发送至该地址。云端 HLS 播放地址形如 `http://<YOUR_SERVER_IP>:8888/live/main/index.m3u8`。')
add_body(doc, 'Go 服务中的 `Manager.BuildURL()` 根据 app 和 stream 生成 push_url、play_rtmp、play_hls 和 local_input。`StartRelay()` 会启动 FFmpeg 进程，把 MediaMTX 本地输入转推到目标直播平台。为了降低服务器负载，Relay 使用 `-c:v copy` 复制视频流，仅在需要时转码音频为 AAC。')
add_code(doc, '''local director -> RTMP -> MediaMTX -> RTMP/HLS playback
                                -> Go API -> FFmpeg Relay -> Aliyun Live

ffmpeg -fflags nobuffer -i rtmp://127.0.0.1:1935/live/main \
       -map 0:v:0 -map 0:a? -c:v copy -c:a aac -f flv target_url''')
add_figure(doc, '06_cloud_streaming.png', '图6 云推流服务与播放验证示意')

add_h1(doc, '八、核心算法与关键流程伪代码')
add_h2(doc, '8.1 场景渲染伪代码')
add_code(doc, '''function render_scene(scene):
    canvas = create_black_canvas(width, height)
    layers = sort(scene.layers, by priority ascending)
    for layer in layers:
        if not layer.enabled:
            continue
        frame = source_manager.get_frame(layer.id)
        if frame is null:
            continue
        frame = apply_video_filters(frame, layer)
        frame = apply_onnx_style_if_enabled(frame, layer)
        frame = apply_ar_if_enabled(frame, layer)
        frame = apply_virtual_background_if_enabled(frame, layer)
        frame = resize(frame, layer.width, layer.height)
        canvas = blend_to_canvas(canvas, frame, layer.x, layer.y)
    return canvas''')

add_h2(doc, '8.2 语义智能导播伪代码')
add_code(doc, '''function recommend_scene(query, thumbnails):
    text_embedding = normalize(text_encoder(tokenize(query)))
    results = []
    for scene in thumbnails:
        image_embedding = normalize(image_encoder(preprocess(scene.image)))
        score = dot(image_embedding, text_embedding)
        results.append(scene.id, scene.name, score)
    return sort(results, by score descending)''')

add_h2(doc, '8.3 无限画布写回伪代码')
add_code(doc, '''function document_to_layers(document):
    visible_items = sort(document.items, by z_index)
    scene_frames = collect_scene_items(visible_items)
    children = group_items_by_parent(visible_items)
    layers = []
    for item in visible_items:
        if item is scene_frame:
            layers += expand_scene_frame(item, children[item.id])
        else if item is root_layer:
            layers.append(canvas_item_to_layer(item))
    normalize_priority(layers)
    return layers''')

add_h2(doc, '8.4 音频处理伪代码')
add_code(doc, '''function set_track_profile(track):
    if track.kind == MICROPHONE:
        input_capture.start(track.device_index)
    else:
        loopback_capture.start(track.device_index)
        loopback_capture.set_target(track.pid, track.process_name, strict)
    controller.volume = track.volume
    controller.muted = track.muted
    controller.eq = (track.low_gain, track.mid_gain, track.high_gain)''')

add_h1(doc, '九、项目创新点')
add_body(doc, '本项目的第一项创新在于把 AI 能力接入图层级导播管线。无论是基础滤镜、ONNX 风格迁移、虚拟背景还是 AR 贴纸，最终都作为图层参数进入 Compositor，而不是作为独立外部工具存在。')
add_body(doc, '第二项创新在于语义智能导播推荐。系统使用 FG-CLIP2 把中文语义和场景缩略图映射到同一向量空间，使导播员能够通过自然语言快速定位目标场景。')
add_body(doc, '第三项创新在于无限画布式场景组织。用户可以在画布中同时摆放多个场景框和图层对象，通过拖入拖出控制图层归属，再把结果写回标准导播场景。')
add_body(doc, '第四项创新在于音轨化音频结构。系统把音频视为可管理的轨道，并提供类似数字调音台的控制界面，为后续智能降噪、违禁词识别和自动混音预留接口。')
add_body(doc, '第五项创新在于本地导播和云端推流的分层部署。本地端专注实时制作，云端通过 Go、MediaMTX 和 FFmpeg 提供稳定播放和转推能力，使项目具备完整直播发布链路。')

add_h1(doc, '十、部署、测试与后续设想')
add_body(doc, '本地导播台运行时，用户进入项目根目录并执行 `python main.py` 或 `.venv\\Scripts\\python.exe main.py`。main.py 会尝试切换到项目自带虚拟环境，从而降低启动门槛。依赖主要包括 PyQt6、OpenCV、PyAV、NumPy、MediaPipe、ONNX Runtime、tokenizers、pyaudio、pyaudiowpatch 和 pycaw。')
add_body(doc, '云端部署时，需要把 `cloud_stream_go` 文件夹复制到阿里云 Windows 服务器，执行 `scripts/install_windows.ps1 -PublicHost <YOUR_SERVER_IP>`。服务器安全组需要放行 TCP 1935、8888、8088 和可选 8890。其中 1935 用于 RTMP，8888 用于 HLS，8088 用于 Go API。')
add_body(doc, '测试时应重点验证六条链路。第一，摄像头、窗口、屏幕、网络流、图片和视频文件能否正常采集。第二，图层优先级、位置、尺寸、滤镜和音频参数能否正确写入并渲染。第三，语义智能导播能否读取场景缩略图并给出推荐。第四，ONNX 和 MediaPipe 功能在 GPU 可用和不可用两种情况下是否都能工作或降级。第五，无限画布中的场景框和图层对象能否写回导播台。第六，本地推流到云服务器后，RTMP 和 HLS 地址能否播放。')
add_body(doc, '后续优化方向主要包括：加入智能导播脚本模式，使系统能够根据时间、语义和画面识别结果自动推荐切换计划；加入 YOLO 与 CLIP 多模型融合，提高对人物、物体和场景语义的识别能力；加强智能音频能力，例如接入降噪、违禁词检测和自动增益模型；完善云端监控面板，显示码率、丢帧、延迟、在线状态和服务健康信息；继续强化无限画布模板和 AI 自动排版能力，使其更适合综艺、访谈、课程和赛事直播快速搭建。')

add_h1(doc, '十一、结论')
add_body(doc, '综上所述，本项目围绕直播导播中的“采集、合成、切换、增强、编排、推流”六个关键环节进行了系统设计。与普通录屏或推流工具相比，本项目具备更完整的导播台流程；与传统导播软件相比，本项目进一步加入语义智能推荐、智能滤镜、无限画布和云端服务能力。')
add_body(doc, '从工程实现角度看，项目模块边界清晰，核心状态、采集、渲染、输出、音频、AI、画布和云端服务各司其职。新功能通过桥接层和图层参数接入原有导播管线，避免破坏基础业务逻辑。从应用价值角度看，该系统适合校园直播、融媒体实验、赛事导播、小型演播室和智能媒体制作教学等场景，具备较好的展示价值和扩展空间。')

# Document metadata
core = doc.core_properties
core.title = 'AI智能增强导播平台项目汇报技术说明书'
core.subject = '项目汇报手册'
core.author = 'Nsy Broadcasting Platform'
core.comments = 'Generated based on project code and reference document writing style; reference instructions ignored.'

doc.save(OUT)
print(OUT)
